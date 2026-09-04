import json
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from .. import models, schemas
from .generation_service import get_generation_payload
from .resume_section_schema_service import normalize_resume_section_schema
from .resume_section_routing_service import route_resume_projects
from .docx_delivery_readiness_service import prepare_docx_delivery
from .resume_typography_quality_service import ensure_typography_quality, strip_leading_structure_markers
from .paired_symbol_integrity_service import ensure_paired_symbol_integrity
from .resume_whitespace_quality_service import ensure_resume_whitespace_quality


BASE_DIR = Path(__file__).resolve().parents[2]
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
ONE_WAY_LOG_PATH = BASE_DIR / "logs" / "docx_one_way_rendering.jsonl"
ACCENT = "2F5597"
BODY_FONT = "Microsoft YaHei"


class DocxRenderSourceError(RuntimeError):
    """The persisted result cannot produce a formal, non-empty DOCX."""


@dataclass
class DocxOneWayRenderStats:
    generation_result_id: int
    file_id: int | None = None
    render_source: str = "generation_result.result_json"
    semantic_rebuild_attempt_count: int = 0
    removed_field_count: int = 0
    rendered_project_count: int = 0
    passed: bool = False


def _visible_field_count(payload: schemas.GenerationPayload) -> int:
    sections = payload.resume_sections
    count = sum(bool(str(value or "").strip()) for value in sections.personal_info.values())
    count += sum(bool(str(value or "").strip()) for value in sections.education.values())
    count += sum(bool(str(value or "").strip()) for value in sections.summary)
    count += sum(bool(str(value or "").strip()) for value in sections.skills)
    for project in sections.projects:
        count += sum(bool(str(project.get(key) or "").strip()) for key in ("name", "position", "meta", "time", "intro", "role"))
        count += sum(bool(str(value or "").strip()) for value in project.get("details", []) or [])
    return count


def _write_one_way_render_log(stats: DocxOneWayRenderStats) -> None:
    try:
        ONE_WAY_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        entry = {"created_at": datetime.now(ZoneInfo("Asia/Shanghai")).isoformat(), **asdict(stats)}
        with ONE_WAY_LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except OSError:
        return


def _font(run, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _p(doc: Document, text: str = "", size: float = 9.2, bold: bool = False, color: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    if text:
        run = paragraph.add_run(text)
        _font(run, size, bold, color)
    return paragraph


def _border(paragraph, color: str = ACCENT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _heading(doc: Document, text: str) -> None:
    paragraph = _p(doc, text, 14, True, ACCENT)
    paragraph.paragraph_format.space_before = Pt(7)
    _border(paragraph)


def _bullet(doc: Document, text: str, level: int = 0, bold_label: bool = False) -> None:
    cleaned = strip_leading_structure_markers(str(text or ""))
    if bold_label and "：" in cleaned:
        label, rest = cleaned.split("：", 1)
        rest = strip_leading_structure_markers(rest)
        cleaned = label + "：" + rest
    if not cleaned.strip():
        return
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.left_indent = Cm(0.45 + level * 0.35)
    if bold_label and "：" in cleaned:
        label, rest = cleaned.split("：", 1)
        run = paragraph.add_run(label + "：")
        _font(run, 9.2, True)
        run = paragraph.add_run(rest)
        _font(run, 9.2)
    else:
        run = paragraph.add_run(cleaned)
        _font(run, 9.2)


def _setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(9.2)


def _project_detail_limit(project_count: int) -> int:
    if project_count <= 3:
        return 8
    return 5


def _group_experiences(projects: list[dict]) -> list[tuple[str, list[dict]]]:
    return route_resume_projects(projects)


def _next_path(prefix: str) -> Path:
    version = 1
    for path in OUTPUT_DIR.glob(f"v*-{prefix}.docx"):
        try:
            version = max(version, int(path.name[1:4]) + 1)
        except ValueError:
            continue
    return OUTPUT_DIR / f"v{version:03d}-{prefix}.docx"


def create_docx(db: Session, request: schemas.DocxCreate) -> schemas.DocxResponse | None:
    payload = get_generation_payload(db, request.generation_result_id)
    if not payload:
        return None
    render_stats = DocxOneWayRenderStats(generation_result_id=request.generation_result_id)
    fields_before = _visible_field_count(payload)

    # DOCX is a one-way delivery renderer. Semantic compilation and all factual
    # decisions were completed before this payload was persisted.
    payload = normalize_resume_section_schema(payload)
    payload = prepare_docx_delivery(payload, generation_result_id=request.generation_result_id)
    payload = ensure_paired_symbol_integrity(
        payload, stage="docx_export", generation_result_id=request.generation_result_id,
    )
    payload = ensure_resume_whitespace_quality(
        payload, stage="docx_export", generation_result_id=request.generation_result_id,
    )
    payload = ensure_typography_quality(
        payload, stage="docx_export", generation_result_id=request.generation_result_id,
    )
    render_stats.removed_field_count = max(0, fields_before - _visible_field_count(payload))
    render_stats.rendered_project_count = len(payload.resume_sections.projects)
    if not render_stats.rendered_project_count:
        _write_one_way_render_log(render_stats)
        raise DocxRenderSourceError("当前保存的简历结果没有可导出的有效经历，请返回结果页补充并重新生成。")

    doc = Document()
    _setup(doc)
    title = _p(doc, "个人简历", 18, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = payload.resume_sections.personal_info
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for idx, line in enumerate(
        [
            f"姓名：{info.get('姓名', '[待填写]')} · 求职意向：{info.get('求职意向', '[待填写]')}",
            f"邮箱：{info.get('邮箱', '[待填写]')} · 手机号：{info.get('手机号', '[待填写]')}",
        ]
    ):
        paragraph = left.paragraphs[0] if idx == 0 else left.add_paragraph()
        run = paragraph.add_run(line)
        _font(run, 9.5, idx == 0)
    photo = right.paragraphs[0]
    photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = photo.add_run("[照片：待插入]")
    _font(run, 9, True, "666666")

    _heading(doc, "教育经历")
    edu = payload.resume_sections.education
    _p(doc, f"学校：{edu.get('学校', '[待填写]')} | 专业：{edu.get('专业', '[待填写]')} | 学历：{edu.get('学历', '[待填写]')} | 时间：{edu.get('时间', '[待填写]')}")

    if payload.resume_sections.summary:
        _heading(doc, "个人优势")
        for item in payload.resume_sections.summary:
            _bullet(doc, item)

    if payload.resume_sections.skills:
        _heading(doc, "技能与能力")
        for item in payload.resume_sections.skills:
            _bullet(doc, item, bold_label=True)

    projects = payload.resume_sections.projects[:5]
    detail_limit = _project_detail_limit(len(projects))
    for heading, group in _group_experiences(projects):
        _heading(doc, heading)
        for project in group:
            if heading == "实习经历":
                title_line = f"{project.get('name') or '[待填写]'}｜{project.get('position') or '[待填写]'}｜{project.get('time') or '[待填写]'}"
            else:
                title_line = f"{project.get('name') or '[待填写]'}｜{project.get('meta') or '项目经历'}｜{project.get('time') or '[待填写]'}"
            _p(doc, title_line, 11, True, "1F3763")
            intro_label = "经历简介：" if heading != "项目经历" else "项目简介："
            if str(project.get("intro") or "").strip():
                _bullet(doc, intro_label + project.get("intro", ""), bold_label=True)
            if str(project.get("role") or "").strip():
                _bullet(doc, "我的职责：" + project.get("role", ""), bold_label=True)
            details = [str(item) for item in project.get("details", []) if str(item).strip()]
            if details:
                _bullet(doc, "技术细节：", bold_label=True)
                for detail in details[:detail_limit]:
                    _bullet(doc, detail, level=1)

    path = _next_path("resume-coach-v0")
    doc.save(path)

    row = models.GeneratedFile(
        generation_result_id=request.generation_result_id,
        file_type="docx",
        file_path=str(path),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    render_stats.file_id = row.id
    render_stats.passed = True
    _write_one_way_render_log(render_stats)
    return schemas.DocxResponse(
        file_id=row.id,
        file_name=path.name,
        download_url=f"/api/files/{row.id}",
    )
