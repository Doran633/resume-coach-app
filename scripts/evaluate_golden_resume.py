import argparse
import json
import os
import re
import sys
import tempfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app import schemas  # noqa: E402
from app.services.experience_fact_ledger_service import build_experience_fact_ledger  # noqa: E402
from app.services.fact_coverage_guard_service import guard_fact_coverage  # noqa: E402
from app.services.experience_boundary_guard_service import guard_experience_boundaries  # noqa: E402
from app.services.resume_section_layering_service import layer_resume_sections  # noqa: E402
from app.services.resume_fact_increment_service import ensure_resume_fact_increment  # noqa: E402
from app.services.resume_fact_dedup_service import deduplicate_resume_facts, similarity  # noqa: E402
from app.services.resume_summary_quality_service import ensure_resume_summary_quality  # noqa: E402
from app.services.resume_skill_evidence_guard_service import guard_resume_skill_evidence  # noqa: E402
from app.services.resume_skill_taxonomy_service import calibrate_resume_skill_taxonomy  # noqa: E402
from app.services.recruiter_facing_technical_language_service import ensure_recruiter_facing_technical_language  # noqa: E402
from app.services.resume_text_integrity_service import ensure_resume_text_integrity  # noqa: E402
from app.services.resume_output_firewall_service import guard_resume_output  # noqa: E402
from app.services.fact_guard_service import guard_hard_facts  # noqa: E402
from app.services.experience_type_resolution_service import resolve_project_types  # noqa: E402
from app.services.resume_title_format_service import resolve_resume_titles  # noqa: E402


DEFAULT_CASES = ROOT / "tests" / "fixtures" / "golden_resume_cases.json"
DEFAULT_OUT = ROOT / "backend" / "reports"


def load_case(case_id: str, path: Path = DEFAULT_CASES) -> dict:
    cases = json.loads(path.read_text(encoding="utf-8"))
    if not case_id:
        return cases[0]
    return next((item for item in cases if item.get("case_id") == case_id), None) or _raise_case(case_id)


def _raise_case(case_id: str):
    raise ValueError(f"Unknown golden case: {case_id}")


def process_fixed_payload(case: dict) -> schemas.GenerationPayload:
    raw = case["raw_input"]
    payload = schemas.GenerationPayload.model_validate(case["fixed_payload"])
    payload = guard_hard_facts(payload, raw)
    payload = guard_fact_coverage(payload, raw, stage="golden_test", write_log=False)
    payload = guard_experience_boundaries(payload, raw, stage="golden_test", write_log=False)
    payload = layer_resume_sections(payload, stage="golden_test", write_log=False)
    payload = ensure_resume_fact_increment(payload)
    payload = deduplicate_resume_facts(payload, stage="golden_test", write_log=False)
    payload = ensure_resume_summary_quality(payload, raw, stage="golden_test", write_log=False)
    payload = guard_resume_skill_evidence(payload, raw, stage="golden_test", write_log=False)
    payload = calibrate_resume_skill_taxonomy(payload, case.get("target_role", ""), stage="golden_test", write_log=False)
    payload = ensure_recruiter_facing_technical_language(payload, stage="golden_test", write_log=False)
    payload = ensure_resume_text_integrity(payload, raw, stage="golden_test", write_log=False)
    payload = guard_resume_output(payload, raw, stage="golden_test", write_log=False)
    payload = guard_hard_facts(payload, raw)
    payload = resolve_project_types(payload, raw, stage="golden_test", write_log=False)
    return resolve_resume_titles(payload, raw)


def visible_text(payload: schemas.GenerationPayload) -> str:
    sections = payload.resume_sections
    parts = [*sections.summary, *sections.skills]
    for project in sections.projects:
        parts.extend(str(project.get(key) or "") for key in ("name", "position", "meta", "time", "intro", "role"))
        parts.extend(str(item) for item in project.get("details", []) or [])
    return "\n".join(parts)


def project_texts(payload: schemas.GenerationPayload) -> dict[str, str]:
    result = {}
    for project in payload.resume_sections.projects:
        source_id = str(project.get("source_experience_id") or "")
        result[source_id] = " ".join([
            str(project.get("name") or ""), str(project.get("position") or ""),
            str(project.get("intro") or ""), str(project.get("role") or ""),
            *[str(item) for item in project.get("details", []) or []],
        ])
    return result


def evaluate_payload(case: dict, payload: schemas.GenerationPayload, docx_text: str = "") -> dict:
    text = visible_text(payload)
    projects = project_texts(payload)
    ledger = build_experience_fact_ledger(case["raw_input"])
    required_results = []
    for fact in case["required_facts"]:
        source_text = projects.get(fact["experience_id"], "")
        covered = all(term.lower() in source_text.lower() for term in fact["all_terms"])
        required_results.append((fact["label"], covered))
    required_covered = sum(covered for _, covered in required_results)

    duplicate_count = 0
    for project in payload.resume_sections.projects:
        rows = [str(item) for item in project.get("details", []) or []]
        duplicate_count += sum(
            similarity(rows[left], rows[right]) >= 0.90
            for left in range(len(rows)) for right in range(left + 1, len(rows))
        )

    internal_leaks = sum(text.lower().count(field.lower()) for field in case["forbidden_internal_fields"])
    expected_categories = case["expected_skill_categories"]
    skills = "\n".join(payload.resume_sections.skills)
    category_checks = [f"{category}：" in skills and all(term in skills for term in terms) for category, terms in expected_categories.items()]
    boundaries = []
    for source_id, rules in case.get("experience_boundaries", {}).items():
        source_text = projects.get(source_id, "")
        boundaries.append(all(term.lower() in source_text.lower() for term in rules.get("required", [])))
        boundaries.append(not any(term.lower() in source_text.lower() for term in rules.get("forbidden", [])))

    actual_types = [str(project.get("resolved_experience_type") or project.get("meta") or "") for project in payload.resume_sections.projects]
    expected_types = case.get("expected_experience_types", [])
    type_matches = sum(
        actual == expected
        for actual, expected in zip(actual_types, expected_types)
    )

    ledger_high = [fact for fact in ledger.facts if fact.importance == "high"]
    return {
        "case_id": case["case_id"],
        "experience_count": len(payload.resume_sections.projects),
        "experience_retention_rate": round(len(payload.resume_sections.projects) / max(1, case["expected_experience_count"]) * 100, 1),
        "explicit_fact_count": len([fact for fact in ledger.facts if fact.explicit]),
        "high_value_fact_count": len(ledger_high),
        "covered_fact_count": required_covered,
        "missing_fact_ids": [label for label, covered in required_results if not covered],
        "fact_coverage_rate": round(required_covered / max(1, len(required_results)) * 100, 1),
        "duplicate_detail_count": duplicate_count,
        "internal_field_leak_count": internal_leaks,
        "skill_category_accuracy": round(sum(category_checks) / max(1, len(category_checks)) * 100, 1),
        "experience_boundary_accuracy": round(sum(boundaries) / max(1, len(boundaries)) * 100, 1),
        "experience_type_accuracy": round(type_matches / max(1, len(expected_types)) * 100, 1),
        "summary_count": len(payload.resume_sections.summary),
        "docx_delivery_ready": bool(docx_text) and "面试准备清单" not in docx_text and not any(field in docx_text for field in case["forbidden_internal_fields"]),
        "forbidden_phrase_hits": [term for term in case["forbidden_phrases"] if term in text],
    }


def _quality_regressions(case: dict, metrics: dict) -> list[str]:
    regressions = []
    if metrics["experience_retention_rate"] < 100:
        regressions.append("主要经历数量减少")
    if metrics["fact_coverage_rate"] < 90:
        regressions.append("高价值事实覆盖率低于 90%")
    if metrics["duplicate_detail_count"]:
        regressions.append("出现重复详情")
    if metrics["internal_field_leak_count"]:
        regressions.append("出现内部字段泄露")
    if metrics["skill_category_accuracy"] < 100:
        regressions.append("技能分类不准确")
    if metrics["experience_boundary_accuracy"] < 100:
        regressions.append("经历事实边界异常")
    if metrics["experience_type_accuracy"] < 100:
        regressions.append("经历类型不准确")
    if not case["summary_min_count"] <= metrics["summary_count"] <= case["summary_max_count"]:
        regressions.append("个人优势数量不符合基线")
    regressions.extend(metrics["forbidden_phrase_hits"])
    return regressions


def _generate_openai(case: dict) -> tuple[schemas.GenerationPayload, str]:
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.database import Base
    from app.services.generation_service import create_generation
    from app.services import docx_service
    from docx import Document

    os.environ["LLM_MODE"] = "openai"
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(bind=engine)
    db = sessionmaker(bind=engine)()
    with tempfile.TemporaryDirectory() as tmpdir:
        old_output = docx_service.OUTPUT_DIR
        try:
            docx_service.OUTPUT_DIR = Path(tmpdir)
            request = schemas.GenerateRequest(
                anonymous_user_id="golden-eval", session_id="golden-eval-session",
                target_role=case["target_role"], mode="full_resume",
                packaging_level=case["packaging_level"], experience_type=case["experience_type"],
                raw_input=case["raw_input"],
            )
            response = create_generation(db, request)
            created = docx_service.create_docx(db, schemas.DocxCreate(
                anonymous_user_id="golden-eval", session_id="golden-eval-session",
                generation_result_id=response.generation_result_id,
            ))
            path = Path(tmpdir) / created.file_name
            docx_text = "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
            return response.result, docx_text
        finally:
            docx_service.OUTPUT_DIR = old_output
            db.close()


def render_fixed_docx(case: dict, payload: schemas.GenerationPayload) -> str:
    """Render the deterministic payload through the production DOCX path."""
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.database import Base
    from app import models
    from app.services import docx_service
    from docx import Document

    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(bind=engine)
    db = sessionmaker(bind=engine)()
    with tempfile.TemporaryDirectory() as tmpdir:
        old_output = docx_service.OUTPUT_DIR
        try:
            docx_service.OUTPUT_DIR = Path(tmpdir)
            experience = models.ExperienceInput(
                anonymous_user_id=None, session_id="golden-eval-session",
                target_role=case["target_role"], mode="full_resume",
                packaging_level=case["packaging_level"], experience_type=case["experience_type"],
                raw_input=case["raw_input"],
            )
            db.add(experience)
            db.flush()
            result = models.GenerationResult(
                experience_input_id=experience.id,
                completeness_score=payload.completeness_score,
                result_json=json.dumps(payload.model_dump(), ensure_ascii=False),
            )
            db.add(result)
            db.commit()
            db.refresh(result)
            created = docx_service.create_docx(db, schemas.DocxCreate(
                anonymous_user_id="golden-eval", session_id="golden-eval-session",
                generation_result_id=result.id,
            ))
            path = Path(tmpdir) / created.file_name
            return "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
        finally:
            docx_service.OUTPUT_DIR = old_output
            db.close()


def write_report(case: dict, metrics: dict, mode: str, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"golden-resume-evaluation-{datetime.now(ZoneInfo('Asia/Shanghai')).date().isoformat()}.md"
    regressions = _quality_regressions(case, metrics)
    lines = [
        "# Golden Resume Evaluation", "", f"- 生成时间：{datetime.now(ZoneInfo('Asia/Shanghai')).isoformat()}",
        f"- 案例：{case['case_id']}", f"- 模式：{mode}", "", "## 质量指标", "",
        f"- 经历保留率：{metrics['experience_retention_rate']}%",
        f"- 高价值事实覆盖率：{metrics['fact_coverage_rate']}%",
        f"- 重复详情数量：{metrics['duplicate_detail_count']}",
        f"- 内部字段泄露数量：{metrics['internal_field_leak_count']}",
        f"- 经历边界准确率：{metrics['experience_boundary_accuracy']}%",
        f"- 经历类型准确率：{metrics['experience_type_accuracy']}%",
        f"- 技能分类准确率：{metrics['skill_category_accuracy']}%",
        f"- 个人优势数量：{metrics['summary_count']}",
        f"- DOCX 投递就绪：{'是' if metrics['docx_delivery_ready'] else '未评估/否'}",
        "", "## 相对 v0.5.7 基线", "",
    ]
    lines.extend([f"- {item}" for item in regressions] if regressions else ["- 未发现质量退化"])
    lines.extend(["", "> 报告只包含聚合指标，不包含用户原始输入、完整推荐版本或完整简历正文。", ""])
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def main() -> int:
    parser = argparse.ArgumentParser(description="Evaluate a golden resume case without exposing source content.")
    parser.add_argument("--case", default="v057_ai_agent_full_resume")
    parser.add_argument("--mode", choices=["mock", "openai"], default="mock")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    case = load_case(args.case)
    if args.mode == "openai":
        payload, docx_text = _generate_openai(case)
    else:
        payload = process_fixed_payload(case)
        docx_text = render_fixed_docx(case, payload)
    metrics = evaluate_payload(case, payload, docx_text)
    report = write_report(case, metrics, args.mode, args.out)
    print(report)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
