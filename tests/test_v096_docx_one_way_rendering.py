import json
import sys
import tempfile
from pathlib import Path

import pytest


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document  # noqa: E402
from sqlalchemy import create_engine  # noqa: E402
from sqlalchemy.orm import sessionmaker  # noqa: E402

from app import models, schemas  # noqa: E402
from app.database import Base  # noqa: E402
from app.services import docx_service  # noqa: E402
from app.services.docx_service import DocxRenderSourceError, create_docx  # noqa: E402


RAW = """项目二：论文阅读助手
使用 React、FastAPI 实现论文检索与摘要阅读功能。

科研说明：此处只用于确认 raw_input 不会影响已经保存的项目类型。"""


def _payload(*, details: list[str] | None = None, name: str = "论文阅读助手", meta: str = "项目经历"):
    return schemas.GenerationPayload(
        completeness_score=80,
        confirmed_facts=[], missing_questions=[], normal_version="", bold_version="", boundary_version="",
        recommended_version="", claims=[], interview_plan=[], knowledge_checklist=[],
        resume_sections=schemas.ResumeSections(
            personal_info={"姓名": "[待填写]"}, education={}, summary=["具备项目实践能力"],
            skills=["React、FastAPI"],
            projects=[{
                "name": name, "meta": meta, "time": "2026.07", "intro": "完成论文检索与阅读功能。",
                "role": "负责前后端功能开发。",
                "details": details or ["使用 React、FastAPI 实现论文检索与摘要阅读功能。"],
                "source_experience_id": "EXP-001",
            }],
        ),
    )


def _db_with_payload(payload: schemas.GenerationPayload, result_id: int = 960):
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(bind=engine)
    db = sessionmaker(bind=engine)()
    db.add(models.ExperienceInput(
        id=1, anonymous_user_id=1, session_id="s", target_role="AI 应用开发",
        mode="full_resume", packaging_level="大胆", experience_type="项目经历", raw_input=RAW,
    ))
    db.add(models.GenerationResult(
        id=result_id, experience_input_id=1, completeness_score=80, result_json=payload.model_dump_json(),
    ))
    db.commit()
    return db


def _block_semantic_rebuild(monkeypatch):
    def fail(*_args, **_kwargs):
        raise AssertionError("DOCX attempted semantic rebuild")

    for name in (
        "build_experience_identities", "resolve_experience_claims", "build_experience_fact_ledger",
        "bind_projects_to_experience_slots", "fill_resume_sections", "guard_fact_coverage",
        "reconcile_resume_projects", "resolve_project_types", "resolve_resume_roles",
    ):
        monkeypatch.setattr(docx_service, name, fail, raising=False)


def test_docx_uses_persisted_project_type_without_reading_raw_input(monkeypatch):
    db = _db_with_payload(_payload())
    _block_semantic_rebuild(monkeypatch)
    old_output = docx_service.OUTPUT_DIR
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            docx_service.OUTPUT_DIR = Path(tmpdir)
            response = create_docx(db, schemas.DocxCreate(
                anonymous_user_id="u", session_id="s", generation_result_id=960,
            ))
            text = "\n".join(item.text for item in Document(Path(tmpdir) / response.file_name).paragraphs)
            assert "项目经历" in text
            assert "科研经历" not in text
            assert "论文阅读助手｜项目经历｜2026.07" in text
        finally:
            docx_service.OUTPUT_DIR = old_output
            db.close()


def test_docx_never_mutates_saved_generation_result_json():
    payload = _payload(details=[" - 使用 React、FastAPI 实现论文检索与摘要阅读功能。"])
    db = _db_with_payload(payload, result_id=961)
    stored_before = db.query(models.GenerationResult).filter_by(id=961).one().result_json
    old_output = docx_service.OUTPUT_DIR
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            docx_service.OUTPUT_DIR = Path(tmpdir)
            response = create_docx(db, schemas.DocxCreate(
                anonymous_user_id="u", session_id="s", generation_result_id=961,
            ))
            text = "\n".join(item.text for item in Document(Path(tmpdir) / response.file_name).paragraphs)
            assert " -" not in text
            assert "使用 React、FastAPI" in text
            assert db.query(models.GenerationResult).filter_by(id=961).one().result_json == stored_before
        finally:
            docx_service.OUTPUT_DIR = old_output
            db.close()


def test_docx_rejects_empty_render_source_without_creating_file():
    db = _db_with_payload(_payload(name="其他经历", meta="项目经历", details=["完成相关工作"]), result_id=962)
    with pytest.raises(DocxRenderSourceError):
        create_docx(db, schemas.DocxCreate(anonymous_user_id="u", session_id="s", generation_result_id=962))
    assert db.query(models.GeneratedFile).filter_by(generation_result_id=962).count() == 0
    db.close()


def test_one_way_render_log_contains_only_aggregate_metadata(tmp_path):
    db = _db_with_payload(_payload(), result_id=963)
    old_output = docx_service.OUTPUT_DIR
    old_log = docx_service.ONE_WAY_LOG_PATH
    try:
        docx_service.OUTPUT_DIR = tmp_path / "outputs"
        docx_service.OUTPUT_DIR.mkdir()
        docx_service.ONE_WAY_LOG_PATH = tmp_path / "docx_one_way_rendering.jsonl"
        create_docx(db, schemas.DocxCreate(anonymous_user_id="u", session_id="s", generation_result_id=963))
        entry = docx_service.ONE_WAY_LOG_PATH.read_text(encoding="utf-8")
        assert RAW not in entry
        assert "论文阅读助手" not in entry
        assert "FastAPI" not in entry
        parsed = json.loads(entry)
        assert parsed["render_source"] == "generation_result.result_json"
        assert parsed["semantic_rebuild_attempt_count"] == 0
        assert parsed["passed"] is True
    finally:
        docx_service.OUTPUT_DIR = old_output
        docx_service.ONE_WAY_LOG_PATH = old_log
        db.close()
